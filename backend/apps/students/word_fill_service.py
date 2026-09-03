"""
Word Fill Service for Salom CRM.

Universal AI-assisted mail-merge engine for university application forms (.docx).

Pipeline:
  1. `analyze_docx_structure` walks the document (tables + paragraphs) and emits an
     addressable slot list -- every place a value could be written, with the label
     text that sits next to it.
  2. `ai_suggest_mapping` sends that skeleton to the configured LLM (same providers
     as ai_ocr_service) and gets back a CRM field suggestion per slot. A regex
     fallback keeps the feature usable when no API key is configured.
  3. The head manager reviews/edits the suggestions in the UI.
  4. `generate_filled_documents` writes CRM values into the approved slots for each
     selected student, producing one .docx per student, preserving all original
     formatting (runs, fonts, borders, merges are never rebuilt -- only run text
     is replaced).
"""

import io
import os
import re
import json
import time
import zipfile
import logging
from copy import deepcopy
from typing import List, Dict, Any, Optional, Tuple

import requests
from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from .excel_fill_service import (
    CRM_FIELDS,
    SEMANTIC_PATTERNS,
    extract_city_and_state,
    format_cell_value,
)

logger = logging.getLogger(__name__)


# ─── Slot kinds ──────────────────────────────────────────────────────────────
# text     : write plain value into an empty cell / after a label
# checkbox : mark one option among several inside a single cell ("M ( ) / F ( )")
# date_parts: fill a "year __ month __ day __" style split date
SLOT_TEXT = 'text'
SLOT_CHECKBOX = 'checkbox'

# Word forms rarely need the system columns Excel does.
WORD_SYSTEM_FIELDS = [
    {
        "key": "_static_value",
        "label": "Statik qiymat (Barchaga bir xil matn)",
        "category": "system",
        "description": "Masalan: Agentlik nomi, E-visa, Round 1",
    },
    {
        "key": "_skip",
        "label": "O'tkazib yuborish (Bo'sh qoldirish)",
        "category": "system",
        "description": "Ushbu joy to'ldirilmaydi",
    },
]

# Extra fields useful in application forms but absent from the Excel list.
WORD_EXTRA_FIELDS = []


def get_word_crm_fields() -> List[Dict[str, Any]]:
    """CRM field dictionary tailored for Word forms (no sequence numbering)."""
    base = [f for f in CRM_FIELDS if f['key'] not in ('_sequence_no', '_static_value', '_skip')]
    seen = set()
    res = []
    for item in (WORD_SYSTEM_FIELDS + base + WORD_EXTRA_FIELDS):
        if item['key'] not in seen:
            seen.add(item['key'])
            res.append(item)
    return res


VALID_FIELD_KEYS = {f['key'] for f in get_word_crm_fields()}


# ─── Placeholder / blank detection ───────────────────────────────────────────

# A cell counts as "fillable" when it is empty or holds only filler characters.
BLANK_PATTERN = re.compile(r'^[\s_\.\-–—…·:•　]*$')

# Existing mail-merge placeholders already present in the template.
PLACEHOLDER_PATTERN = re.compile(r'\{\{\s*([a-zA-Z0-9_]+)\s*\}\}')

# Checkbox-ish option groups: "M (  ) / F (  )", "남( ) 여( )", "Yes[ ] No[ ]"
CHECKBOX_OPTION_PATTERN = re.compile(
    r'([^\s()\[\]/·,]{1,24}?)\s*[\(\[]\s*[\)\]]',
    re.UNICODE,
)

# Cells that are pure instruction/example text should not be offered as slots.
NOISE_PREFIXES = ('※', '▪', '▣', '*', '·')


def _is_blank_text(text: str) -> bool:
    return bool(BLANK_PATTERN.match(text or ''))


def _looks_like_label(text: str) -> bool:
    """A label is short-ish descriptive text, not a paragraph of instructions."""
    t = (text or '').strip()
    if not t or len(t) > 90:
        return False
    return True


def _detect_checkbox_options(text: str) -> List[Dict[str, str]]:
    """
    Finds selectable options inside a cell, e.g. "M (   ) / F (   )".
    Returns [{"label": "M", "match": "M (   )"}, ...] when at least 2 are present.
    """
    options = []
    for m in CHECKBOX_OPTION_PATTERN.finditer(text or ''):
        label = m.group(1).strip()
        if not label or label.startswith(NOISE_PREFIXES):
            continue
        options.append({"label": label, "span": [m.start(), m.end()]})
    return options if len(options) >= 2 else []


# ─── Document walking ────────────────────────────────────────────────────────

def _iter_block_items(parent):
    """Yields Paragraph and Table objects in true document order."""
    if isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent.element.body

    for child in parent_elm.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)


def _cell_key(cell: _Cell) -> str:
    """Identity of the underlying <w:tc> element -- merged cells share one."""
    return str(id(cell._tc))


def _collect_table_slots(
    table: Table,
    table_path: str,
    slots: List[Dict[str, Any]],
) -> None:
    """
    Scans one table, pairing label cells with the blank cell that follows them.

    University forms are almost always "label | value" grids with heavy merging,
    so the value cell is the first non-duplicate blank cell to the right of a
    label (falling back to the cell directly below for vertical layouts).
    """
    seen_in_table: set = set()
    # (row, label) pairs already given a slot -- stops one label from claiming a
    # whole run of narrow per-character boxes.
    claimed_labels: set = set()
    grid: List[List[Tuple[str, _Cell]]] = []

    for r_idx, row in enumerate(table.rows):
        row_cells: List[Tuple[str, _Cell]] = []
        for c_idx, cell in enumerate(row.cells):
            row_cells.append((_cell_key(cell), cell))
        grid.append(row_cells)

    for r_idx, row_cells in enumerate(grid):
        prev_key = None
        for c_idx, (key, cell) in enumerate(row_cells):
            # Skip horizontally merged duplicates
            if key == prev_key:
                continue
            prev_key = key

            text = cell.text.strip()

            # 1. Existing {{placeholder}} in the template -> direct slot
            ph = PLACEHOLDER_PATTERN.search(text)
            if ph:
                slot_id = f"{table_path}.r{r_idx}.c{c_idx}"
                if slot_id in seen_in_table:
                    continue
                seen_in_table.add(slot_id)
                slots.append({
                    "slot_id": slot_id,
                    "kind": SLOT_TEXT,
                    "label": _nearest_label(grid, r_idx, c_idx),
                    "current_text": text,
                    "existing_placeholder": ph.group(1),
                    "table_index": int(table_path.split('t')[-1]),
                    "row": r_idx,
                    "col": c_idx,
                })
                continue

            # 2. Checkbox groups ("M ( ) / F ( )")
            options = _detect_checkbox_options(text)
            if options:
                slot_id = f"{table_path}.r{r_idx}.c{c_idx}"
                if slot_id in seen_in_table:
                    continue
                seen_in_table.add(slot_id)
                slots.append({
                    "slot_id": slot_id,
                    "kind": SLOT_CHECKBOX,
                    "label": _nearest_label(grid, r_idx, c_idx),
                    "current_text": text,
                    "options": [o["label"] for o in options],
                    "table_index": int(table_path.split('t')[-1]),
                    "row": r_idx,
                    "col": c_idx,
                })
                continue

            # 3. Blank cell that has a label to its left (or above).
            #    Forms often split one value across several narrow boxes (one per
            #    character); only the first blank after a given label is offered.
            if _is_blank_text(text):
                label = _nearest_label(grid, r_idx, c_idx)
                if not label:
                    continue

                label_marker = (r_idx, label)
                if label_marker in claimed_labels:
                    continue
                claimed_labels.add(label_marker)

                slot_id = f"{table_path}.r{r_idx}.c{c_idx}"
                if slot_id in seen_in_table:
                    continue
                seen_in_table.add(slot_id)
                slots.append({
                    "slot_id": slot_id,
                    "kind": SLOT_TEXT,
                    "label": label,
                    "current_text": "",
                    "table_index": int(table_path.split('t')[-1]),
                    "row": r_idx,
                    "col": c_idx,
                })


def _nearest_label(
    grid: List[List[Tuple[str, _Cell]]],
    r_idx: int,
    c_idx: int,
) -> str:
    """
    Finds the descriptive text that identifies this cell: nearest non-blank cell
    to the left on the same row, else the nearest non-blank cell above.
    Merged duplicates are skipped so a wide header does not shadow the real label.
    """
    row = grid[r_idx]
    own_key = row[c_idx][0]

    # Left neighbours (closest first)
    for c in range(c_idx - 1, -1, -1):
        key, cell = row[c]
        if key == own_key:
            continue
        text = cell.text.strip()
        if text and not _is_blank_text(text) and _looks_like_label(text):
            return text

    # Above neighbours
    for r in range(r_idx - 1, -1, -1):
        if c_idx >= len(grid[r]):
            continue
        key, cell = grid[r][c_idx]
        if key == own_key:
            continue
        text = cell.text.strip()
        if text and not _is_blank_text(text) and _looks_like_label(text):
            return text

    return ""


def _collect_paragraph_slots(doc, slots: List[Dict[str, Any]]) -> None:
    """
    Picks up {{placeholder}} markers and "Label: ______" lines outside tables.
    """
    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        ph = PLACEHOLDER_PATTERN.search(text)
        if ph:
            slots.append({
                "slot_id": f"p{p_idx}",
                "kind": SLOT_TEXT,
                "label": re.sub(PLACEHOLDER_PATTERN, '', text).strip() or text,
                "current_text": text,
                "existing_placeholder": ph.group(1),
                "paragraph_index": p_idx,
            })
            continue

        # "Name: ________" / "Date ______"
        m = re.match(r'^(.{2,60}?)\s*[:：]?\s*[_\.]{3,}\s*$', text)
        if m:
            slots.append({
                "slot_id": f"p{p_idx}",
                "kind": SLOT_TEXT,
                "label": m.group(1).strip(),
                "current_text": text,
                "paragraph_index": p_idx,
            })


def analyze_docx_structure(file_bytes: bytes) -> Dict[str, Any]:
    """
    Builds the addressable slot list for a .docx template.
    Returns slots plus lightweight table previews for the review UI.
    """
    doc = Document(io.BytesIO(file_bytes))
    slots: List[Dict[str, Any]] = []

    for t_idx, table in enumerate(doc.tables):
        _collect_table_slots(table, f"t{t_idx}", slots)

    _collect_paragraph_slots(doc, slots)

    # Table previews (deduplicated cells) so the head manager sees context.
    tables_preview = []
    for t_idx, table in enumerate(doc.tables):
        rows_preview = []
        for r_idx, row in enumerate(table.rows[:25]):
            prev_key = None
            values = []
            for cell in row.cells:
                key = _cell_key(cell)
                if key == prev_key:
                    continue
                prev_key = key
                values.append(cell.text.strip()[:60])
            if any(values):
                rows_preview.append({"row_idx": r_idx, "values": values})
        tables_preview.append({
            "table_index": t_idx,
            "rows": len(table.rows),
            "cols": len(table.columns),
            "preview_rows": rows_preview,
        })

    return {
        "slots": slots,
        "tables": tables_preview,
        "paragraph_count": len(doc.paragraphs),
        "available_fields": get_word_crm_fields(),
    }


# ─── Regex fallback matching (no AI key configured) ──────────────────────────

def _regex_match_field(label: str) -> Tuple[str, float]:
    """Reuses the Excel semantic dictionary against a Word label."""
    if not label or not str(label).strip():
        return ('_skip', 0.0)

    cleaned = " ".join(str(label).strip().split()).lower()
    for pattern, field_key in SEMANTIC_PATTERNS:
        if re.search(pattern, cleaned, re.IGNORECASE):
            if field_key in VALID_FIELD_KEYS:
                return (field_key, 0.8)
    return ('_skip', 0.1)


# Checkbox groups only make sense for fields whose value picks one option.
CHECKBOX_SAFE_FIELDS = {'gender', 'level', 'nationality'}


def _fallback_mapping(slots: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    out = []
    for slot in slots:
        existing = slot.get('existing_placeholder')
        if existing and existing in VALID_FIELD_KEYS:
            out.append({
                "slot_id": slot['slot_id'],
                "field": existing,
                "confidence": 1.0,
                "reason": "Shablonda mavjud placeholder",
            })
            continue

        field, conf = _regex_match_field(slot.get('label', ''))

        # A label like "Applying Major" sitting over a campus picker would
        # otherwise tick an unrelated box; leave those for the reviewer.
        if slot.get('kind') == SLOT_CHECKBOX and field not in CHECKBOX_SAFE_FIELDS:
            out.append({
                "slot_id": slot['slot_id'],
                "field": "_skip",
                "confidence": 0.0,
                "reason": "Variantli maydon — qo'lda tanlang",
            })
            continue

        out.append({
            "slot_id": slot['slot_id'],
            "field": field,
            "confidence": conf,
            "reason": "Lug'at asosida moslandi" if field != '_skip' else "Mos maydon topilmadi",
        })
    return out


# ─── AI suggestion ───────────────────────────────────────────────────────────

def _build_ai_prompt(slots: List[Dict[str, Any]]) -> str:
    field_lines = "\n".join(
        f"- {f['key']}: {f['label']}" for f in get_word_crm_fields()
    )

    slot_lines = []
    for s in slots:
        kind = s['kind']
        desc = f"slot_id={s['slot_id']} | kind={kind} | label=\"{s.get('label', '')}\""
        if kind == SLOT_CHECKBOX:
            desc += f" | options={s.get('options', [])}"
        if s.get('current_text'):
            desc += f" | current_text=\"{s['current_text'][:80]}\""
        slot_lines.append(desc)

    return f"""You are mapping blank fields of a Korean/English university application form (.docx) to CRM database fields.

AVAILABLE CRM FIELDS:
{field_lines}

FORM SLOTS TO MAP:
{chr(10).join(slot_lines)}

RULES:
1. For each slot_id decide which CRM field belongs there, based on its label.
2. Labels may be Korean, English, Uzbek or Russian. "성명"=name, "여권번호"=passport, "생년월일"=birthday, "국적"=nationality, "연락처"=phone, "주소"=address, "부"=father, "모"=mother.
3. "[Kor]" next to a name label means korean_name; "[Eng]" means full_name.
4. kind=checkbox slots are option groups (e.g. M/F). Map them only when a CRM field decides the choice (usually gender). Otherwise "_skip".
5. If nothing fits, or the slot is a signature/essay/free-text area, use "_skip".
6. confidence is 0.0-1.0 reflecting how certain the mapping is.

Return ONLY valid JSON in exactly this shape:
{{"mappings": [{{"slot_id": "t0.r2.c3", "field": "full_name", "confidence": 0.95, "reason": "short reason"}}]}}"""


def _call_llm(prompt: str, provider: str, api_key: Optional[str], model: Optional[str]) -> str:
    """Mirrors the provider handling of ai_ocr_service, text-only."""
    provider_clean = (provider or 'openai').lower().strip()

    if provider_clean == 'gemini':
        key = api_key or os.environ.get('GEMINI_API_KEY') or os.environ.get('GOOGLE_API_KEY')
        if not key:
            raise Exception("Gemini API Key sozlanmagan")

        models_to_try = [model or 'gemini-2.5-flash']
        if 'gemini-1.5-flash' not in models_to_try:
            models_to_try.append('gemini-1.5-flash')

        payload = {
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {"response_mime_type": "application/json"},
        }

        last_error = None
        for current_model in models_to_try:
            url = (
                f"https://generativelanguage.googleapis.com/v1beta/models/"
                f"{current_model}:generateContent?key={key}"
            )
            for _ in range(2):
                try:
                    resp = requests.post(
                        url, json=payload,
                        headers={"Content-Type": "application/json"},
                        timeout=90,
                    )
                    if resp.status_code in (429, 503):
                        time.sleep(1.0)
                        continue
                    resp.raise_for_status()
                    candidates = resp.json().get("candidates", [])
                    if candidates and "content" in candidates[0]:
                        parts = candidates[0]["content"].get("parts", [])
                        if parts:
                            return parts[0].get("text", "{}")
                except Exception as e:
                    last_error = e
                    time.sleep(0.5)
        raise Exception(f"Gemini so'rovi muvaffaqiyatsiz: {last_error}")

    key = api_key or os.environ.get('OPENAI_API_KEY')
    if not key:
        raise Exception("OpenAI API Key sozlanmagan")

    payload = {
        "model": model or 'gpt-4o',
        "response_format": {"type": "json_object"},
        "messages": [{"role": "user", "content": prompt}],
    }
    resp = requests.post(
        "https://api.openai.com/v1/chat/completions",
        headers={"Content-Type": "application/json", "Authorization": f"Bearer {key}"},
        json=payload,
        timeout=90,
    )
    resp.raise_for_status()
    return resp.json().get("choices", [{}])[0].get("message", {}).get("content", "{}")


def _parse_ai_json(raw: str) -> List[Dict[str, Any]]:
    text = (raw or '').strip()
    if text.startswith('```'):
        text = re.sub(r'^```(?:json)?\s*', '', text)
        text = re.sub(r'\s*```$', '', text)
    try:
        data = json.loads(text)
    except Exception:
        m = re.search(r'\{.*\}', text, re.DOTALL)
        if not m:
            raise
        data = json.loads(m.group(0))
    return data.get('mappings', []) if isinstance(data, dict) else []


def ai_suggest_mapping(
    slots: List[Dict[str, Any]],
    provider: str = 'openai',
    api_key: Optional[str] = None,
    model: Optional[str] = None,
) -> Tuple[List[Dict[str, Any]], str]:
    """
    Asks the LLM to map each slot to a CRM field.
    Returns (mappings, source) where source is "ai", "ai_partial" or "fallback".
    Never raises -- degrades to the regex dictionary so the page always works.
    """
    if not slots:
        return ([], 'fallback')

    fallback = _fallback_mapping(slots)

    try:
        raw = _call_llm(_build_ai_prompt(slots), provider, api_key, model)
        ai_maps = _parse_ai_json(raw)
    except Exception as e:
        logger.warning(f"Word Fill AI mapping unavailable, using dictionary fallback: {e}")
        return (fallback, 'fallback')

    by_id = {m['slot_id']: m for m in ai_maps if isinstance(m, dict) and m.get('slot_id')}
    valid_slot_ids = {s['slot_id'] for s in slots}

    merged = []
    ai_hits = 0
    for fb in fallback:
        sid = fb['slot_id']
        cand = by_id.get(sid)
        if cand and sid in valid_slot_ids:
            field = cand.get('field', '_skip')
            if field in VALID_FIELD_KEYS:
                ai_hits += 1
                try:
                    conf = float(cand.get('confidence', 0.7))
                except (TypeError, ValueError):
                    conf = 0.7
                merged.append({
                    "slot_id": sid,
                    "field": field,
                    "confidence": max(0.0, min(1.0, conf)),
                    "reason": str(cand.get('reason', ''))[:160],
                })
                continue
        merged.append(fb)

    source = 'ai' if ai_hits >= len(slots) * 0.5 else 'ai_partial'
    return (merged, source)


# ─── Value resolution ────────────────────────────────────────────────────────

def resolve_field_value(
    field_key: str,
    student: Dict[str, Any],
    static_value: str = '',
    fallback: str = '',
    format_rules: Optional[Dict[str, Any]] = None,
) -> str:
    """Turns a CRM field key into the string to write for one student."""
    rules = format_rules or {}

    if field_key == '_static_value':
        return static_value or ''
    if field_key == 'today_date':
        from datetime import datetime as _dt
        return _dt.now().strftime('%Y-%m-%d')
    if field_key in ('first_name', 'last_name'):
        parts = (student.get('full_name') or '').split()
        if field_key == 'first_name':
            return parts[0] if parts else fallback
        return " ".join(parts[1:]) if len(parts) > 1 else fallback
    if field_key == 'nationality':
        return student.get('nationality') or 'UZBEKISTAN'
    if field_key == 'address_city':
        city, _ = extract_city_and_state(student.get('address') or '')
        return city or fallback
    if field_key == 'address_state':
        _, state = extract_city_and_state(student.get('address') or '')
        return state or fallback

    return format_cell_value(student.get(field_key), field_key, rules, fallback)


# ─── Writing into the document ───────────────────────────────────────────────

def _clone_run_formatting(source_run, target_run) -> None:
    """Copies the <w:rPr> block so a new run inherits font, size, colour and weight."""
    if source_run is None or target_run is None:
        return
    src_rPr = source_run._element.rPr
    if src_rPr is None:
        return
    tgt = target_run._element
    if tgt.rPr is not None:
        tgt.remove(tgt.rPr)
    tgt.insert(0, deepcopy(src_rPr))


def _usable_donor(para) -> Optional[Any]:
    """
    Returns the best run in a paragraph to copy formatting from: one that both
    holds text and carries an explicit <w:rPr>. A text-bearing run without rPr is
    accepted as a weaker second choice, so plainly styled templates still donate
    their paragraph-level look instead of falling through to Word's defaults.
    """
    weak = None
    for run in para.runs:
        if not run.text.strip():
            continue
        if run._element.rPr is not None:
            return run
        if weak is None:
            weak = run
    return weak


def _find_style_donor_run(cell: _Cell, table: Optional[Table] = None):
    """
    Finds a run whose formatting a brand-new run should copy.

    An empty form cell usually has no runs at all, so writing into it with a bare
    add_run() would produce default-styled text that clashes with the rest of the
    form. Preference order: a formatted run already in this cell, then the label
    cells nearest to it, which in these grids carry the form's real font.
    """
    for para in cell.paragraphs:
        donor = _usable_donor(para)
        if donor is not None:
            return donor

    if table is None:
        return None

    own_key = _cell_key(cell)

    # Locate the cell's row so the search can radiate outwards from it.
    own_row_idx = None
    for r_idx, row in enumerate(table.rows):
        if any(_cell_key(c) == own_key for c in row.cells):
            own_row_idx = r_idx
            break

    row_order = range(len(table.rows))
    if own_row_idx is not None:
        row_order = sorted(
            range(len(table.rows)),
            key=lambda r: abs(r - own_row_idx),
        )

    for r_idx in row_order:
        for other in table.rows[r_idx].cells:
            if _cell_key(other) == own_key:
                continue
            for para in other.paragraphs:
                donor = _usable_donor(para)
                if donor is not None:
                    return donor
    return None


def _set_cell_text(cell: _Cell, text: str, style_donor=None) -> None:
    """
    Replaces a cell's text while keeping its formatting.

    The first run of the first paragraph carries the cell's styling, so the value
    goes there and every other run is emptied -- this preserves font, size, colour
    and cell borders exactly as the university designed them. When the cell is
    truly empty a new run is created and given the donor's formatting.
    """
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.text = text
        return

    first_para = paragraphs[0]
    if first_para.runs:
        first_para.runs[0].text = text
        for run in first_para.runs[1:]:
            run.text = ''
    else:
        new_run = first_para.add_run(text)
        _clone_run_formatting(style_donor, new_run)

    for para in paragraphs[1:]:
        for run in para.runs:
            run.text = ''


def _replace_in_paragraph(para: Paragraph, old: str, new: str) -> bool:
    """
    Replaces text inside a paragraph even when Word split it across runs.
    Formatting of the first affected run is kept.
    """
    full = para.text
    if old not in full:
        return False

    start = full.index(old)
    end = start + len(old)

    pos = 0
    first_hit = None
    for run in para.runs:
        run_len = len(run.text)
        run_start, run_end = pos, pos + run_len
        pos = run_end

        if run_end <= start or run_start >= end:
            continue

        local_start = max(0, start - run_start)
        local_end = min(run_len, end - run_start)

        if first_hit is None:
            first_hit = run
            run.text = run.text[:local_start] + new + run.text[local_end:]
        else:
            run.text = run.text[:local_start] + run.text[local_end:]

    return first_hit is not None


def _apply_checkbox(cell: _Cell, options: List[str], chosen: str, mark: str = 'V') -> None:
    """
    Marks the option matching `chosen` inside a "M ( ) / F ( )" style cell by
    putting a mark between that option's brackets, leaving the rest untouched.
    """
    if not chosen:
        return

    chosen_norm = chosen.strip().upper()
    if not chosen_norm:
        return

    normalized = [(opt, opt.strip().upper()) for opt in options]

    # Exact match, then prefix either way ("M" vs "MALE"), never a bare first
    # letter -- that would let FEMALE select the M box.
    target = next((o for o, n in normalized if n == chosen_norm), None)
    if target is None:
        target = next(
            (o for o, n in normalized if n and (n.startswith(chosen_norm) or chosen_norm.startswith(n))),
            None,
        )
    if target is None:
        return

    for para in cell.paragraphs:
        text = para.text
        pattern = re.compile(
            re.escape(target) + r'\s*([\(\[])(\s*)([\)\]])'
        )
        m = pattern.search(text)
        if not m:
            continue
        old = m.group(0)
        new = old.replace(
            f"{m.group(1)}{m.group(2)}{m.group(3)}",
            f"{m.group(1)}{mark}{m.group(3)}",
        )
        if _replace_in_paragraph(para, old, new):
            return


def _resolve_slot_targets(doc) -> Dict[str, Tuple[Any, Optional[Table]]]:
    """
    Maps slot_id -> (live object, owning table) in this document instance.
    The table is carried along so an empty cell can borrow formatting from its row.
    """
    targets: Dict[str, Tuple[Any, Optional[Table]]] = {}

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            prev_key = None
            for c_idx, cell in enumerate(row.cells):
                key = _cell_key(cell)
                if key == prev_key:
                    continue
                prev_key = key
                targets[f"t{t_idx}.r{r_idx}.c{c_idx}"] = (cell, table)

    for p_idx, para in enumerate(doc.paragraphs):
        targets[f"p{p_idx}"] = (para, None)

    return targets


def fill_single_document(
    file_bytes: bytes,
    mappings: List[Dict[str, Any]],
    student: Dict[str, Any],
    checkbox_mark: str = 'V',
) -> io.BytesIO:
    """Produces one filled .docx for one student."""
    doc = Document(io.BytesIO(file_bytes))
    targets = _resolve_slot_targets(doc)

    for mapping in mappings:
        field_key = mapping.get('field', '_skip')
        if not field_key or field_key == '_skip':
            continue

        slot_id = mapping.get('slot_id')
        if not slot_id:
            continue
        entry = targets.get(str(slot_id))
        if entry is None:
            continue
        target, owning_table = entry

        value = resolve_field_value(
            field_key,
            student,
            static_value=mapping.get('static_value', ''),
            fallback=mapping.get('fallback', ''),
            format_rules=mapping.get('format_rules', {}),
        )

        kind = mapping.get('kind', SLOT_TEXT)

        if kind == SLOT_CHECKBOX and isinstance(target, _Cell):
            _apply_checkbox(target, mapping.get('options', []), value, checkbox_mark)
            continue

        existing_ph = mapping.get('existing_placeholder')

        if isinstance(target, _Cell):
            donor = _find_style_donor_run(target, owning_table)
            if existing_ph:
                replaced = False
                for para in target.paragraphs:
                    if _replace_in_paragraph(para, '{{' + existing_ph + '}}', value):
                        replaced = True
                        break
                if not replaced:
                    _set_cell_text(target, value, donor)
            else:
                _set_cell_text(target, value, donor)
        elif isinstance(target, Paragraph):
            if existing_ph:
                _replace_in_paragraph(target, '{{' + existing_ph + '}}', value)
            else:
                label = mapping.get('label', '')
                m = re.search(r'[_\.]{3,}', target.text)
                if m:
                    _replace_in_paragraph(target, m.group(0), value)
                elif label:
                    _replace_in_paragraph(target, target.text, f"{label}: {value}")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _safe_filename(name: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|]+', '_', str(name or '')).strip()
    cleaned = re.sub(r'\s+', ' ', cleaned)
    return cleaned[:120] or 'document'


def build_output_filename(pattern: str, student: Dict[str, Any], index: int) -> str:
    """
    Renders a filename from a pattern supporting {full_name}, {id}, {passport},
    {index} and {date} tokens.
    """
    from datetime import datetime as _dt

    tokens = {
        'full_name': student.get('full_name') or 'Student',
        'id': student.get('id') or '',
        'passport': student.get('passport') or '',
        'index': str(index),
        'date': _dt.now().strftime('%Y-%m-%d'),
    }

    result = pattern or '{full_name}'
    for key, val in tokens.items():
        result = result.replace('{' + key + '}', str(val))

    result = _safe_filename(result)
    if not result.lower().endswith('.docx'):
        result += '.docx'
    return result


def generate_filled_documents(
    file_bytes: bytes,
    mappings: List[Dict[str, Any]],
    students_data: List[Dict[str, Any]],
    filename_pattern: str = '{full_name}',
    checkbox_mark: str = 'V',
) -> Tuple[io.BytesIO, str, int]:
    """
    Fills the template for every selected student.

    One student  -> a single .docx stream.
    Many students -> a .zip archive of individually named .docx files.
    Returns (stream, kind, count) where kind is "docx" or "zip".
    """
    if not students_data:
        raise ValueError("Hech bo'lmaganda bitta talaba tanlanishi kerak")

    active = [m for m in mappings if m.get('field') and m.get('field') != '_skip']
    if not active:
        raise ValueError("Hech bo'lmaganda bitta maydon moslanishi kerak")

    if len(students_data) == 1:
        stream = fill_single_document(file_bytes, active, students_data[0], checkbox_mark)
        return (stream, 'docx', 1)

    zip_buffer = io.BytesIO()
    used_names: Dict[str, int] = {}

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for idx, student in enumerate(students_data, start=1):
            doc_stream = fill_single_document(file_bytes, active, student, checkbox_mark)

            name = build_output_filename(filename_pattern, student, idx)
            if name in used_names:
                used_names[name] += 1
                stem = name[:-5]
                name = f"{stem} ({used_names[name]}).docx"
            else:
                used_names[name] = 1

            zf.writestr(name, doc_stream.getvalue())

    zip_buffer.seek(0)
    return (zip_buffer, 'zip', len(students_data))
