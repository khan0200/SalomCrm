"""
Tests for the Word Fill engine.

Templates are built in-memory with python-docx so the suite never depends on the
university .docx files kept outside version control.
"""

import io
import zipfile
import unittest

from docx import Document

from apps.students.word_fill_service import (
    analyze_docx_structure,
    _fallback_mapping,
    _detect_checkbox_options,
    fill_single_document,
    generate_filled_documents,
    build_output_filename,
    resolve_field_value,
    get_word_crm_fields,
)


def _horizontal_form() -> bytes:
    """Label in one column, blank value cell to its right."""
    doc = Document()
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = 'Full Name'
    table.cell(1, 0).text = 'Passport Number'
    table.cell(2, 0).text = 'Date of Birth'
    table.cell(3, 0).text = 'Gender'
    table.cell(3, 1).text = 'M (   ) / F (   )'

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _vertical_form() -> bytes:
    """Labels on the header row, blanks directly underneath."""
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = 'English Name'
    table.cell(0, 1).text = 'Nationality'
    table.cell(0, 2).text = 'e-mail Address'

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _placeholder_form() -> bytes:
    """Template already carrying {{mail-merge}} markers."""
    doc = Document()
    doc.add_paragraph('Applicant: {{full_name}}')
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = 'Passport'
    table.cell(0, 1).text = '{{passport}}'

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


STUDENT = {
    'id': '1024',
    'full_name': 'ABDUVOIDOV KHAYITALI',
    'korean_name': '압두보이도프',
    'passport': 'AD1234567',
    'birthday': '2003-05-14',
    'gender': 'MALE',
    'nationality': 'UZBEKISTAN',
    'email': 'khayitali@mail.uz',
    'phone1': '901234567',
    'address': 'ANDIJAN REGION, ASAKA DISTRICT',
}


def _mappings_for(template: bytes):
    """Runs analysis + dictionary mapping and merges slot metadata, as the UI does."""
    res = analyze_docx_structure(template)
    maps = _fallback_mapping(res['slots'])
    slot_by_id = {s['slot_id']: s for s in res['slots']}
    for m in maps:
        slot = slot_by_id[m['slot_id']]
        m['kind'] = slot['kind']
        m['options'] = slot.get('options', [])
        m['existing_placeholder'] = slot.get('existing_placeholder')
        m['format_rules'] = {}
    return res, maps


def _all_cell_texts(doc) -> list:
    texts = []
    for table in doc.tables:
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                key = id(cell._tc)
                if key in seen:
                    continue
                seen.add(key)
                texts.append(cell.text.strip())
    return texts


class TestCheckboxDetection(unittest.TestCase):
    def test_detects_slash_separated_options(self):
        opts = _detect_checkbox_options('M (   ) / F (   )')
        self.assertEqual([o['label'] for o in opts], ['M', 'F'])

    def test_detects_korean_options(self):
        opts = _detect_checkbox_options('남( ) 여( )')
        self.assertEqual([o['label'] for o in opts], ['남', '여'])

    def test_single_option_is_not_a_checkbox_group(self):
        self.assertEqual(_detect_checkbox_options('Signature (    )'), [])

    def test_plain_text_has_no_options(self):
        self.assertEqual(_detect_checkbox_options('Full Name'), [])


class TestStructureAnalysis(unittest.TestCase):
    def test_finds_labels_to_the_left(self):
        res = analyze_docx_structure(_horizontal_form())
        labels = {s['label'] for s in res['slots']}
        self.assertIn('Full Name', labels)
        self.assertIn('Passport Number', labels)
        self.assertIn('Date of Birth', labels)

    def test_finds_labels_above(self):
        res = analyze_docx_structure(_vertical_form())
        labels = {s['label'] for s in res['slots']}
        self.assertEqual(labels, {'English Name', 'Nationality', 'e-mail Address'})

    def test_checkbox_slot_carries_options(self):
        res = analyze_docx_structure(_horizontal_form())
        checkbox = [s for s in res['slots'] if s['kind'] == 'checkbox']
        self.assertEqual(len(checkbox), 1)
        self.assertEqual(checkbox[0]['options'], ['M', 'F'])

    def test_existing_placeholders_are_detected(self):
        res = analyze_docx_structure(_placeholder_form())
        found = {s.get('existing_placeholder') for s in res['slots']}
        self.assertIn('full_name', found)
        self.assertIn('passport', found)

    def test_available_fields_exclude_sequence_numbering(self):
        keys = {f['key'] for f in get_word_crm_fields()}
        self.assertNotIn('_sequence_no', keys)
        self.assertIn('_skip', keys)
        self.assertIn('full_name', keys)


class TestDictionaryMapping(unittest.TestCase):
    def test_maps_common_labels(self):
        res, maps = _mappings_for(_horizontal_form())
        by_label = {
            next(s['label'] for s in res['slots'] if s['slot_id'] == m['slot_id']): m['field']
            for m in maps
        }
        self.assertEqual(by_label['Full Name'], 'full_name')
        self.assertEqual(by_label['Passport Number'], 'passport')
        self.assertEqual(by_label['Date of Birth'], 'birthday')
        self.assertEqual(by_label['Gender'], 'gender')

    def test_gender_checkbox_is_mapped(self):
        _, maps = _mappings_for(_horizontal_form())
        checkbox = [m for m in maps if m['kind'] == 'checkbox']
        self.assertEqual(len(checkbox), 1)
        self.assertEqual(checkbox[0]['field'], 'gender')

    def test_unrelated_checkbox_groups_are_left_for_review(self):
        """A campus/branch picker must not be auto-ticked from a 'Major' label."""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = 'Applying Major'
        table.cell(0, 1).text = 'Anyang (   ) / Ganghwa (   )'
        buf = io.BytesIO()
        doc.save(buf)

        _, maps = _mappings_for(buf.getvalue())
        checkbox = [m for m in maps if m['kind'] == 'checkbox']
        self.assertEqual(len(checkbox), 1)
        self.assertEqual(checkbox[0]['field'], '_skip')

    def test_existing_placeholder_wins_with_full_confidence(self):
        _, maps = _mappings_for(_placeholder_form())
        placeholder_maps = [m for m in maps if m.get('existing_placeholder')]
        self.assertTrue(placeholder_maps)
        for m in placeholder_maps:
            self.assertEqual(m['confidence'], 1.0)
            self.assertEqual(m['field'], m['existing_placeholder'])


class TestFilling(unittest.TestCase):
    def test_writes_values_into_blank_cells(self):
        template = _horizontal_form()
        _, maps = _mappings_for(template)
        out = fill_single_document(template, maps, STUDENT)
        texts = _all_cell_texts(Document(out))
        self.assertIn('ABDUVOIDOV KHAYITALI', texts)
        self.assertIn('AD1234567', texts)
        self.assertIn('2003-05-14', texts)

    def test_marks_only_the_matching_checkbox_option(self):
        template = _horizontal_form()
        _, maps = _mappings_for(template)
        out = fill_single_document(template, maps, STUDENT)
        gender_cell = Document(out).tables[0].cell(3, 1).text
        self.assertIn('M (V)', gender_cell)
        self.assertIn('F (   )', gender_cell)

    def test_female_marks_the_second_option(self):
        template = _horizontal_form()
        _, maps = _mappings_for(template)
        female = dict(STUDENT, gender='FEMALE')
        out = fill_single_document(template, maps, female)
        gender_cell = Document(out).tables[0].cell(3, 1).text
        self.assertIn('F (V)', gender_cell)
        self.assertNotIn('M (V)', gender_cell)

    def test_label_cells_are_never_overwritten(self):
        template = _horizontal_form()
        _, maps = _mappings_for(template)
        out = fill_single_document(template, maps, STUDENT)
        table = Document(out).tables[0]
        self.assertEqual(table.cell(0, 0).text, 'Full Name')
        self.assertEqual(table.cell(1, 0).text, 'Passport Number')

    def test_replaces_existing_placeholders(self):
        template = _placeholder_form()
        _, maps = _mappings_for(template)
        out = fill_single_document(template, maps, STUDENT)
        doc = Document(out)
        self.assertIn('ABDUVOIDOV KHAYITALI', doc.paragraphs[0].text)
        self.assertNotIn('{{', doc.paragraphs[0].text)
        self.assertEqual(doc.tables[0].cell(0, 1).text, 'AD1234567')

    def test_skipped_fields_leave_cells_untouched(self):
        template = _horizontal_form()
        _, maps = _mappings_for(template)
        for m in maps:
            m['field'] = '_skip'
        out = fill_single_document(template, maps, STUDENT)
        self.assertEqual(Document(out).tables[0].cell(0, 1).text, '')

    def test_value_is_written_as_a_single_run(self):
        """One run per value keeps the cell's styling coherent in Word."""
        template = _horizontal_form()
        _, maps = _mappings_for(template)
        out = fill_single_document(template, maps, STUDENT)
        cell = Document(out).tables[0].cell(0, 1)
        self.assertEqual(len(cell.paragraphs[0].runs), 1)
        self.assertEqual(cell.paragraphs[0].runs[0].text, 'ABDUVOIDOV KHAYITALI')

    def test_new_runs_inherit_explicit_run_formatting(self):
        """When the template styles its runs, a filled value copies that styling."""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        label_run = table.cell(0, 0).paragraphs[0].add_run('Full Name')
        label_run.font.name = 'Batang'
        label_run.font.bold = True

        buf = io.BytesIO()
        doc.save(buf)
        template = buf.getvalue()

        _, maps = _mappings_for(template)
        out = fill_single_document(template, maps, STUDENT)
        run = Document(out).tables[0].cell(0, 1).paragraphs[0].runs[0]
        self.assertEqual(run.text, 'ABDUVOIDOV KHAYITALI')
        self.assertEqual(run.font.name, 'Batang')
        self.assertTrue(run.font.bold)

    def test_original_layout_is_preserved(self):
        template = _horizontal_form()
        _, maps = _mappings_for(template)
        out = fill_single_document(template, maps, STUDENT)
        original, filled = Document(io.BytesIO(template)), Document(out)
        self.assertEqual(len(original.tables), len(filled.tables))
        self.assertEqual(len(original.tables[0].rows), len(filled.tables[0].rows))
        self.assertEqual(len(original.tables[0].columns), len(filled.tables[0].columns))


class TestBatchGeneration(unittest.TestCase):
    def test_single_student_returns_one_docx(self):
        template = _horizontal_form()
        _, maps = _mappings_for(template)
        stream, kind, count = generate_filled_documents(template, maps, [STUDENT])
        self.assertEqual((kind, count), ('docx', 1))
        self.assertIn('ABDUVOIDOV KHAYITALI', _all_cell_texts(Document(stream)))

    def test_many_students_return_a_zip_of_documents(self):
        template = _horizontal_form()
        _, maps = _mappings_for(template)
        students = [
            dict(STUDENT, id='1', full_name='ALI VALIYEV'),
            dict(STUDENT, id='2', full_name='ZUHRA KARIMOVA', gender='FEMALE'),
        ]
        stream, kind, count = generate_filled_documents(
            template, maps, students, filename_pattern='{full_name}_{id}'
        )
        self.assertEqual((kind, count), ('zip', 2))

        archive = zipfile.ZipFile(stream)
        self.assertEqual(
            sorted(archive.namelist()),
            ['ALI VALIYEV_1.docx', 'ZUHRA KARIMOVA_2.docx'],
        )
        first = Document(io.BytesIO(archive.read('ALI VALIYEV_1.docx')))
        self.assertIn('ALI VALIYEV', _all_cell_texts(first))

    def test_ten_students_yield_ten_named_files_in_one_zip(self):
        """Each selected student gets their own .docx, named after them, in a single archive."""
        template = _horizontal_form()
        _, maps = _mappings_for(template)

        names = [
            'ABDUVOIDOV KHAYITALI', 'KAMBAROV SIROJIDDIN', 'KARIMOVA ZUHRA',
            'TOSHMATOV BEKZOD', 'YULDASHEVA MOHIRA', 'RAHIMOV JAVOHIR',
            'SAIDOVA NILUFAR', 'ERGASHEV OYBEK', 'NORMATOVA DILNOZA',
            'QODIROV SANJAR',
        ]
        students = [
            dict(STUDENT, id=str(1000 + i), full_name=n, passport=f'AD{1000 + i}')
            for i, n in enumerate(names, start=1)
        ]

        stream, kind, count = generate_filled_documents(
            template, maps, students, filename_pattern='{full_name}'
        )
        self.assertEqual((kind, count), ('zip', 10))

        archive = zipfile.ZipFile(stream)
        self.assertEqual(sorted(archive.namelist()), sorted(f'{n}.docx' for n in names))

        # Every document must carry its own student's data and nobody else's.
        for student in students:
            entry = f"{student['full_name']}.docx"
            texts = _all_cell_texts(Document(io.BytesIO(archive.read(entry))))
            self.assertIn(student['full_name'], texts, f'{entry} missing its own name')
            self.assertIn(student['passport'], texts, f'{entry} missing its own passport')

            for other in students:
                if other is not student:
                    self.assertNotIn(other['full_name'], texts, f'{entry} leaked another student')

    def test_duplicate_filenames_are_disambiguated(self):
        template = _horizontal_form()
        _, maps = _mappings_for(template)
        twins = [
            dict(STUDENT, id='1', full_name='ALI VALIYEV'),
            dict(STUDENT, id='2', full_name='ALI VALIYEV'),
        ]
        stream, _, _ = generate_filled_documents(
            template, maps, twins, filename_pattern='{full_name}'
        )
        names = zipfile.ZipFile(stream).namelist()
        self.assertEqual(len(names), len(set(names)))

    def test_rejects_empty_student_list(self):
        template = _horizontal_form()
        _, maps = _mappings_for(template)
        with self.assertRaises(ValueError):
            generate_filled_documents(template, maps, [])

    def test_rejects_when_every_field_is_skipped(self):
        template = _horizontal_form()
        _, maps = _mappings_for(template)
        for m in maps:
            m['field'] = '_skip'
        with self.assertRaises(ValueError):
            generate_filled_documents(template, maps, [STUDENT])


class TestFilenameBuilding(unittest.TestCase):
    def test_expands_tokens(self):
        name = build_output_filename('{id}_{full_name}', STUDENT, 3)
        self.assertEqual(name, '1024_ABDUVOIDOV KHAYITALI.docx')

    def test_strips_characters_windows_rejects(self):
        student = dict(STUDENT, full_name='A/B:C*D?E')
        self.assertNotRegex(build_output_filename('{full_name}', student, 1), r'[\\/:*?"<>|]')

    def test_falls_back_when_pattern_renders_empty(self):
        self.assertTrue(build_output_filename('', dict(STUDENT, full_name=''), 1).endswith('.docx'))


class TestValueResolution(unittest.TestCase):
    def test_splits_full_name(self):
        self.assertEqual(resolve_field_value('first_name', STUDENT), 'ABDUVOIDOV')
        self.assertEqual(resolve_field_value('last_name', STUDENT), 'KHAYITALI')

    def test_derives_city_and_region_from_address(self):
        self.assertEqual(resolve_field_value('address_city', STUDENT), 'ASAKA')
        self.assertEqual(resolve_field_value('address_state', STUDENT), 'ANDIJAN')

    def test_nationality_defaults_to_uzbekistan(self):
        self.assertEqual(resolve_field_value('nationality', {'nationality': None}), 'UZBEKISTAN')

    def test_static_value_is_written_verbatim(self):
        self.assertEqual(
            resolve_field_value('_static_value', STUDENT, static_value='E-VISA'),
            'E-VISA',
        )

    def test_missing_value_uses_the_fallback(self):
        self.assertEqual(resolve_field_value('email', {}, fallback='-'), '-')

    def test_date_format_rule_is_applied(self):
        self.assertEqual(
            resolve_field_value('birthday', STUDENT, format_rules={'dateFormat': 'DD.MM.YYYY'}),
            '14.05.2003',
        )


class TestGenderFormatting(unittest.TestCase):
    """
    FEMALE contains an 'M', so a substring test silently wrote MALE for every
    female student. These cases pin the corrected behaviour.
    """

    def _gender(self, raw, fmt='MALE/FEMALE'):
        return resolve_field_value('gender', {'gender': raw}, format_rules={'genderFormat': fmt})

    def test_female_is_not_read_as_male(self):
        for raw in ('FEMALE', 'Female', 'female', 'F'):
            self.assertEqual(self._gender(raw), 'FEMALE', f"misread {raw!r}")

    def test_male_stays_male(self):
        for raw in ('MALE', 'Male', 'M'):
            self.assertEqual(self._gender(raw), 'MALE', f"misread {raw!r}")

    def test_uzbek_words(self):
        self.assertEqual(self._gender('AYOL'), 'FEMALE')
        self.assertEqual(self._gender('ERKAK'), 'MALE')

    def test_korean_output_format(self):
        self.assertEqual(self._gender('FEMALE', '남/여'), '여')
        self.assertEqual(self._gender('MALE', '남/여'), '남')

    def test_short_output_format(self):
        self.assertEqual(self._gender('FEMALE', 'M/F'), 'F')
        self.assertEqual(self._gender('MALE', 'M/F'), 'M')


if __name__ == '__main__':
    unittest.main()
